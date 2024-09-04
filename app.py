from flask import Flask, render_template, request, send_file
from io import BytesIO
from docx import Document
from forms import LegalDocumentForm

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your_secret_key'  # Replace with your own secret key

# Predefined templates for each agreement type
agreement_templates = {
    'partnership': 'This Partnership Agreement is made between {party_1} and {party_2}. The partners agree to share profits and losses equally. Specific Terms: {specific_terms}.',
    'lease': 'This Lease Agreement is made between {party_1} (Landlord) and {party_2} (Tenant). The landlord and the tenant agree on the rental rate and the duration of the lease. Specific Terms: {specific_terms}.',
    'employment': 'This Employment Agreement is made between {party_1} (Employer) and {party_2} (Employee). The employer and the employee agree on the terms of employment, including salary, benefits, and duties. Specific Terms: {specific_terms}.',
    'nda': 'This Non-Disclosure Agreement (NDA) is made between {party_1} and {party_2}. The parties agree to keep the shared information confidential. Specific Terms: {specific_terms}.',
    'licensing': 'This Licensing Agreement is made between {party_1} (Licensor) and {party_2} (Licensee). The licensor agrees to grant the licensee the right to use its intellectual property. Specific Terms: {specific_terms}.',
    'service': 'This Service Agreement is made between {party_1} (Service Provider) and {party_2} (Client). The service provider and the client agree on the scope and price of the services. Specific Terms: {specific_terms}.',
    'settlement': 'This Settlement Agreement is made between {party_1} and {party_2}. The parties agree to settle the dispute and release each other from any further claims. Specific Terms: {specific_terms}.',
    'purchase': 'This Purchase Agreement is made between {party_1} (Seller) and {party_2} (Buyer). The buyer and the seller agree on the terms of the sale, including the price and the delivery date. Specific Terms: {specific_terms}.',
}

@app.route('/', methods=['GET', 'POST'])
def home():
    form = LegalDocumentForm()
    if form.validate_on_submit():
        # Extract data from the form
        agreement_type = form.agreement_type.data
        party_1 = form.party_1.data
        party_2 = form.party_2.data
        specific_terms = form.specific_terms.data
        
        # Get the template and format it with the user input
        agreement_text = agreement_templates[agreement_type].format(
            party_1=party_1, party_2=party_2, specific_terms=specific_terms)
        
        # Create a new Word document
        doc = Document()
        doc.add_heading('Legal Agreement', 0)
        doc.add_paragraph(agreement_text)
        
        # Save the document to a BytesIO object
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Send the document as a download
        return send_file(buffer, as_attachment=True, download_name='legal_agreement.docx')
    
    return render_template('home.html', form=form)

if __name__ == '__main__':
    app.run(debug=True)
